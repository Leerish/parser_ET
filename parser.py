import mimetypes

def detect_type(path):
    mime = mimetypes.guess_type(path)[0]
    print("Detected MIME:", mime)
    return mime


import pypdfium2 as pdfium

def parse_pdf(path):
    print("Parsing PDF with PDFium...")
    doc = pdfium.PdfDocument(path)
    text = []
    for page in doc:
        txt = page.get_textpage().get_text_range()
        text.append(txt)
    return "\n".join(text)


from docx import Document

def parse_docx(path):
    print("Parsing DOCX file...")
    doc = Document(path)
    return "\n".join(p.text for p in doc.paragraphs)


from paddleocr import PaddleOCR
ocr = PaddleOCR(use_angle_cls=True, lang='en')

def parse_image(path):
    print("OCR on image file...")
    result = ocr.ocr(path, cls=True)
    lines = []
    for block in result:
        for line in block:
            lines.append(line[1][0])
    return "\n".join(lines)

def parse_resume(path):
    mime = detect_type(path)

    if mime == "application/pdf":
        return parse_pdf(path)

    elif mime in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
        return parse_docx(path)

    elif mime and mime.startswith("image/"):
        return parse_image(path)

    else:
        raise ValueError("Unsupported file format")

path = "/content/data.pdf"

text = parse_resume(path)

print("========== EXTRACTED TEXT ==========")
print(text)

import re

def extract_email(text):
    pattern = r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}"
    match = re.search(pattern, text)
    return match.group(0) if match else None


def extract_linkedin(text):
    pattern = r"(https?://(www\.)?linkedin\.com/[A-Za-z0-9_/?\-=%.]+)"
    match = re.search(pattern, text)
    return match.group(0) if match else None



def extract_name(text):
    lines = text.strip().split("\n")

    
    for line in lines:
        clean = line.strip()
        
        if "@" in clean or "linkedin.com" in clean.lower():
            continue
        
        if re.match(r"^[A-Za-z ,.\-]+$", clean) and len(clean.split()) <= 5:
            return clean
    return None


SKILL_LIST = [
    "python","java","c++","c","sql","nosql","mysql","postgresql","mongodb","aws",
    "azure","gcp","docker","kubernetes","pytorch","tensorflow","scikit-learn",
    "react","django","flask","fastapi","git","linux","javascript","html","css",
    "nlp","machine learning","deep learning","computer vision","tableau","power bi"
]

def extract_skills(text):
    text_low = text.lower()
    skills_found = []

    for skill in SKILL_LIST:
        if skill in text_low:
            skills_found.append(skill)

    return sorted(list(set(skills_found)))



def extract_work_experience(text):
    patterns = [
        r"(work experience|experience|professional experience|employment)([\s\S]+?)(education|projects|skills|certifications|summary|$)",
    ]

    text_low = text.lower()

    for pattern in patterns:
        match = re.search(pattern, text_low)
        if match:
            return match.group(2).strip()

    return None



def extract_rest(text, work_exp):
    if not work_exp:
        return text

    idx = text.lower().find(work_exp.lower())
    if idx == -1:
        return text

    return text[idx + len(work_exp):].strip()


def parse_information(text):
    info = {}

    info["name"] = extract_name(text)
    info["email"] = extract_email(text)
    info["linkedin"] = extract_name(text)
    info["skills"] = extract_skills(text)

    work_exp = extract_work_experience(text)
    info["work_experience"] = work_exp
    info["rest"] = extract_rest(text, work_exp)

    return info

parsed = parse_information(text)

parsed