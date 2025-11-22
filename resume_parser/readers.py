import mimetypes
import pypdfium2 as pdfium
from docx import Document
from paddleocr import PaddleOCR

import logging
logging.getLogger("ppocr").setLevel(logging.ERROR)
logging.getLogger("paddle").setLevel(logging.ERROR)


# Create OCR instance once (expensive to load repeatedly)
ocr = PaddleOCR(use_angle_cls=True, lang='en')


def detect_type(path: str) -> str | None:
    """
    Detect MIME type of a file by path.
    Returns None if detection fails.
    """
    mime = mimetypes.guess_type(path)[0]
    print("Detected MIME:", mime)
    return mime


def parse_pdf(path: str) -> str:
    print("Parsing PDF with PDFium...")
    doc = pdfium.PdfDocument(path)
    text_chunks = []
    for page in doc:
        txt = page.get_textpage().get_text_range()
        text_chunks.append(txt)
    return "\n".join(text_chunks)


def parse_docx(path: str) -> str:
    print("Parsing DOCX file...")
    doc = Document(path)
    return "\n".join(p.text for p in doc.paragraphs)


def parse_image(path: str) -> str:
    print("OCR on image file...")
    result = ocr.ocr(path, cls=True)
    lines = []
    for block in result:
        for line in block:
            # line[1][0] is the recognized text
            lines.append(line[1][0])
    return "\n".join(lines)


def parse_resume(path: str) -> str:
    """
    High-level dispatcher that picks the correct parser based on MIME type.
    """
    mime = detect_type(path)

    if mime == "application/pdf":
        return parse_pdf(path)

    elif mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return parse_docx(path)

    elif mime and mime.startswith("image/"):
        return parse_image(path)

    else:
        raise ValueError(f"Unsupported file format: {mime}")
